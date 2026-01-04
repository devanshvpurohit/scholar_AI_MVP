
"""
Scholar AI Backend
Refactored to pure Flask for simplicity and flexibility.
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import json
import logging
import tempfile
import time
from werkzeug.utils import secure_filename
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from google.oauth2 import service_account
from google.cloud import speech, storage
import firebase_admin
from firebase_admin import credentials, auth
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pypdf
from mutagen.mp3 import MP3
from mutagen.wave import WAVE
from mutagen.mp4 import MP4

# Initialize Flask App
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================================================
# CONFIGURATION
# ============================================================================

# Initialize Firebase Admin SDK
logger.info("Starting Firebase Init...")
try:
    cred_path = os.environ.get('GOOGLE_APPLICATION_CREDENTIALS', 'service-account-key.json')
    if not firebase_admin._apps:
        if os.path.exists(cred_path):
            cred = credentials.Certificate(cred_path)
            firebase_admin.initialize_app(cred)
            logger.info("Initialized Firebase with service account")
        else:
            # Check if running in a cloud environment where default creds work
            try:
                # Explicitly set project ID for local development with user credentials
                # firebase_admin.initialize_app(options={'projectId': 'medkey-vault'})
                # logger.info("Initialized Firebase Admin with project: medkey-vault")
                logger.warning("Skipping Firebase ADC Init to prevent crash - Use Demo User")
            except Exception as e:
                logger.warning(f"Firebase Admin init failed: {e}")
except Exception as e:
    logger.error(f"Critical Firebase Init Error: {e}")
logger.info("Firebase Init Step Complete")

# Configure Gemini API
# Keys should be placed in backend-functions/.env file
DEFAULT_GEMINI_KEY = os.environ.get("GEMINI_API_KEY")

def get_gemini_key(request=None):
    """Get Gemini API key from request header or fallback"""
    if request:
        header_key = request.headers.get('X-Gemini-API-Key')
        if header_key and header_key.strip():
            return header_key
    
    if DEFAULT_GEMINI_KEY and DEFAULT_GEMINI_KEY != "PASTE_YOUR_GEMINI_API_KEY_HERE":
        return DEFAULT_GEMINI_KEY
    return None

def configure_genai(request=None):
    """Configure GenAI with specific key"""
    api_key = get_gemini_key(request)
    if api_key:
        genai.configure(api_key=api_key)
        return True
    return False

# Google Cloud Storage bucket
GCS_BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME", "scholar-ai-storage")

# Temporary storage for guides (In-memory/File-based for demo)
# In production, use Firestore or a real database
DB_PATH = '/tmp/scholar_ai_db'
if not os.path.exists(DB_PATH):
    os.makedirs(DB_PATH)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def verify_firebase_token(req):
    """Verify Firebase ID token from Authorization header"""
    auth_header = req.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        # Allow demo access if explicitly requested or handling local dev issues
        if os.environ.get('FLASK_DEBUG') == '1' or app.debug:
            logger.warning("Auth header missing in generic dev mode. Defaulting to demo_user.")
            return {'uid': 'demo_user'}, None
        return None, "Missing or invalid Authorization header"
    
    token = auth_header.split('Bearer ')[1]
    
    # Bypass verification if Firebase is not initialized (Local Dev Mode)
    if not firebase_admin._apps:
        logger.warning("Firebase not initialized. Returning demo_user for local dev.")
        return {'uid': 'demo_user', 'email': 'demo@scholar.ai'}, None

    try:
        decoded_token = auth.verify_id_token(token)
        return decoded_token, None
    except Exception as e:
        logger.error(f"Token verification failed: {e}")
        # If credentials are missing locally or app not init, fallback to demo user
        error_msg = str(e).lower()
        if "credentials" in error_msg or "app does not exist" in error_msg:
             logger.warning("Firebase Auth Error (Dev). Falling back to demo_user.")
             return {'uid': 'demo_user'}, None
        return None, str(e)

def upload_blob(bucket_name, source_file_name, destination_blob_name):
    """Uploads a file to the bucket."""
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)
    blob.upload_from_filename(source_file_name)
    return f"gs://{bucket_name}/{destination_blob_name}"

def delete_blob(bucket_name, blob_name):
    """Deletes a blob from the bucket."""
    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        blob.delete()
    except Exception as e:
        logger.warning(f"Failed to delete blob {blob_name}: {e}")

def get_audio_duration(file_path):
    """Get duration of audio file in seconds"""
    try:
        if file_path.endswith('.mp3'):
            audio = MP3(file_path)
        elif file_path.endswith('.wav'):
            audio = WAVE(file_path)
        elif file_path.endswith('.mp4'): # Treat video audio
            audio = MP4(file_path)
        else:
            return 0
        return audio.info.length
    except Exception as e:
        logger.error(f"Error getting audio duration: {e}")
        return 0

def get_audio_transcript(file_path):
    """Transcribe audio using Google Cloud Speech-to-Text"""
    client = speech.SpeechClient()
    
    duration = get_audio_duration(file_path)
    logger.info(f"Audio duration: {duration} seconds")

    # Use long running recognize for longer files (> 55s to be safe)
    if duration > 55:
        logger.info("Using Long Running Recognize via GCS")
        bucket_name = GCS_BUCKET_NAME
        blob_name = f"temp_audio_{int(time.time())}_{os.path.basename(file_path)}"
        gcs_uri = upload_blob(bucket_name, file_path, blob_name)
        
        audio = speech.RecognitionAudio(uri=gcs_uri)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16 if file_path.endswith('.wav') else speech.RecognitionConfig.AudioEncoding.MP3,
            language_code="en-US",
            enable_automatic_punctuation=True,
        )
        
        operation = client.long_running_recognize(config=config, audio=audio)
        response = operation.result(timeout=600)
        
        # Clean up GCS
        delete_blob(bucket_name, blob_name)
        
    else:
        # Synchronous recognize for short files
        with open(file_path, "rb") as audio_file:
            content = audio_file.read()
            
        audio = speech.RecognitionAudio(content=content)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16 if file_path.endswith('.wav') else speech.RecognitionConfig.AudioEncoding.MP3,
            language_code="en-US",
            enable_automatic_punctuation=True,
        )
        response = client.recognize(config=config, audio=audio)

    transcript = ""
    for result in response.results:
        transcript += result.alternatives[0].transcript + " "
        
    return transcript.strip()

# ============================================================================
# GEMINI PROMPTS (Updated to gemini-1.5-flash)
# ============================================================================

def prompt_everything(prompt, goals=""):
    """Generate EVERYTHING in a single Gemini call to save time (crucial for Vercel timeout)"""
    model = genai.GenerativeModel("gemini-2.5-flash-lite")
    
    final_prompt = (
        f"You are an expert, proactive study planner. The user's specific goal is: '{goals}'. "
        "Analyze the transcript and generate a comprehensive study system. "
        "1. **Tag Topics by Difficulty**: Identify key topics and rate them (Easy/Medium/Hard). "
        "2. **Spaced Repetition**: Insert specific 'Revision' slots in the schedule for hard topics to ensure retention. "
        "3. **Feasibility**: Ensure daily study time is realistic (30-60 mins max per session). "
        "4. **Personalization**: Provide specific study tips for this material. "
        "Return a SINGLE VALID JSON object with this EXACT structure:\n"
        "{\n"
        "  \"title\": \"Catchy Title\",\n"
        "  \"summary\": \"Detailed summary...\",\n"
        "  \"topics\": [\n"
        "     {\"name\": \"Topic A\", \"difficulty\": \"Hard\"},\n"
        "     {\"name\": \"Topic B\", \"difficulty\": \"Medium\"}\n"
        "  ],\n"
        "  \"study_tips\": [\"Tip 1\", \"Tip 2 using mnemonic...\"],\n"
        "  \"flash_cards\": [[\"Q1\", \"A1\"], ...], (10 cards)\n"
        "  \"quiz\": [\n"
        "    {\"question\": \"Q1\", \"possible_answers\": [\"A\",\"B\",\"C\",\"D\"], \"index\": 0, \"related_topic\": \"Topic A\"},\n"
        "    ... (10 questions)\n"
        "  ],\n"
        "  \"study_schedule\": [\n"
        "     {\"day_offset\": 1, \"title\": \"Study: Topic A\", \"details\": \"Deep dive...\", \"duration_minutes\": 45, \"type\": \"learning\", \"difficulty\": \"Hard\"},\n"
        "     {\"day_offset\": 2, \"title\": \"Revision: Topic A\", \"details\": \"Quick review to reinforce memory.\", \"duration_minutes\": 15, \"type\": \"revision\", \"difficulty\": \"Hard\"},\n"
        "     {\"day_offset\": 3, \"title\": \"Quiz & Assessment\", \"details\": \"Final check.\", \"duration_minutes\": 20, \"type\": \"quiz\", \"difficulty\": \"Medium\"}\n"
        "  ]\n"
        "}\n\n"
        "Transcript:\n" + prompt
    )
    
    try:
        response = model.generate_content(final_prompt, generation_config={"response_mime_type": "application/json"})
        return json.loads(response.text)
    except Exception as e:
        logger.error(f"Gemini Generation Error: {e}")
        # Fallback to empty structure if parsing fails
        return {
            "title": "Error Generating Guide",
            "summary": "The AI could not process this text within the time limit or format constraints.",
            "flash_cards": [],
            "quiz": [],
            "study_schedule": []
        }

# Deprecated individual prompt functions removed for performance
# def prompt_flashcards...
# def prompt_quiz...
# def prompt_summary...
# def prompt_title...

def clean_json_response(text):
    """Clean markdown code blocks from JSON response"""
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:]
    if text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return json.loads(text)

# ============================================================================
# API ROUTES
# ============================================================================

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy", "service": "Scholar AI Backend"}), 200

@app.route('/api/models', methods=['GET'])
def list_available_models():
    """List available Gemini models"""
    configure_genai(request)
    try:
        models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                models.append(m.name)
        return jsonify({"models": models})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def handle_upload():
    # Verify Auth
    user, error = verify_firebase_token(request)
    if not user:
        return jsonify({"error": f"Unauthorized: {error}"}), 401
    user_id = user['uid'] 
    
    # Configure Gemini with user provided key
    if not configure_genai(request):
        return jsonify({"error": "Gemini API Key is missing. Please provide it in the input field."}), 400

    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    try:
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        
        logger.info(f"Processing file: {filename}")
        
        transcript = ""
        
        # 1. Process File to get Text
        if filename.endswith('.pdf'):
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                transcript += page.extract_text() + "\n"
                
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            doc = Document(file_path)
            for para in doc.paragraphs:
                transcript += para.text + "\n"
                
        elif filename.endswith('.txt') or filename.endswith('.md'):
            with open(file_path, 'r') as f:
                transcript = f.read()
                
        elif filename.endswith('.mp3') or filename.endswith('.wav') or filename.endswith('.mp4'):
            # For audio, we can either use Speech-to-Text OR use Gemini 1.5 Pro/Flash directly
            # Gemini 1.5 Flash supports audio input directly!
            # For now, let's keep the STT path as fallback or specific use case, 
            # OR we could upload the file to Gemini API directly.
            # Sticking to STT logic for now to preserve existing business logic, 
            # but Gemini 1.5 Flash is multimodal.
            
            transcript = get_audio_transcript(file_path)
        
        else:
            return jsonify({"error": "Unsupported file type"}), 400

        if not transcript.strip():
            return jsonify({"error": "Could not extract text from file"}), 400

        # 2. Generate Content
        goals = request.form.get('goals', '')
        study_guide = prompt_everything(transcript, goals)
        
        if study_guide.get('title') == "Error Generating Guide":
            return jsonify({"error": f"AI Generation Failed. Please check your API Key. (Summary: {study_guide.get('summary')})"}), 500
        
        # 3. Store Result
        guide_id = str(int(time.time())) # Simple ID gen
        study_guide['id'] = guide_id
        study_guide['created_at'] = int(time.time())
        study_guide['filename'] = filename
        study_guide['user_id'] = user_id
        study_guide['goals'] = goals
        
        # Save to local JSON "DB"
        with open(os.path.join(DB_PATH, f'{guide_id}.json'), 'w') as f:
            json.dump(study_guide, f)
            
        return jsonify(study_guide), 200

    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/guides', methods=['GET'])
def get_all_guides():
    user, error = verify_firebase_token(request)
    if not user:
        return jsonify({"error": f"Unauthorized: {error}"}), 401
    user_id = user['uid']

    try:
        guides = []
        if os.path.exists(DB_PATH):
            for f in os.listdir(DB_PATH):
                if f.endswith('.json'):
                    with open(os.path.join(DB_PATH, f), 'r') as file:
                        data = json.load(file)
                        # Filter by user
                        if data.get('user_id') == user_id:
                            guides.append({
                                'id': data['id'],
                                'title': data['title'],
                                'filename': data['filename'],
                                'created_at': data['created_at']
                            })
        
        # Sort by newest first
        guides.sort(key=lambda x: x['created_at'], reverse=True)
        return jsonify({"guides": guides}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/guide/<guide_id>', methods=['GET'])
def get_guide(guide_id):
    try:
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        if not os.path.exists(file_path):
            return jsonify({"error": "Guide not found"}), 404
            
        with open(file_path, 'r') as f:
            data = json.load(f)
            
        return jsonify(data), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/guide/<guide_id>', methods=['DELETE'])
def delete_guide(guide_id):
    try:
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        if os.path.exists(file_path):
            os.remove(file_path)
            return jsonify({"success": True}), 200
        return jsonify({"error": "Guide not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Export routes (Flashcards DOCX, Quiz DOCX, Summary DOCX)
@app.route('/api/export/flashcards/<guide_id>', methods=['GET'])
def export_flashcards(guide_id):
    try:
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        with open(file_path, 'r') as f:
            data = json.load(f)
            
        doc = Document()
        doc.add_heading('Flashcards: ' + data['title'], 0)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Answer'

        for card in data['flash_cards']:
            row_cells = table.add_row().cells
            row_cells[0].text = card[0]
            row_cells[1].text = card[1]
            
            # Add simple shading (Simulating "Reddish" / "Orangeish")
            # Note: Complex XML manipulation needed for exact colors, simplified here
            
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return send_file(
                tmp.name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'flashcards_{guide_id}.docx'
            )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/export/quiz/<guide_id>', methods=['GET'])
def export_quiz(guide_id):
    try:
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        with open(file_path, 'r') as f:
            data = json.load(f)
            
        doc = Document()
        doc.add_heading('Quiz: ' + data['title'], 0)

        for i, q in enumerate(data['quiz']):
            doc.add_paragraph(f"{i+1}. {q['question']}", style='List Number')
            for j, ans in enumerate(q['possible_answers']):
                doc.add_paragraph(f"   {chr(65+j)}. {ans}")
            doc.add_paragraph("") # Space

        # Answer Key
        doc.add_page_break()
        doc.add_heading('Answer Key', 1)
        for i, q in enumerate(data['quiz']):
            correct_ans = q['possible_answers'][q['index']]
            doc.add_paragraph(f"{i+1}. {chr(65+q['index'])} - {correct_ans}")

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return send_file(
                tmp.name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'quiz_{guide_id}.docx'
            )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/export/summary/<guide_id>', methods=['GET'])
def export_summary(guide_id):
    try:
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        with open(file_path, 'r') as f:
            data = json.load(f)
            
        doc = Document()
        doc.add_heading('Summary: ' + data['title'], 0)
        doc.add_paragraph(data['summary']) # Naive markdown dump

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return send_file(
                tmp.name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'summary_{guide_id}.docx'
            )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/guide/<guide_id>/progress', methods=['PUT'])
def update_progress(guide_id):
    try:
        data = request.json
        session_index = data.get('index')
        completed = data.get('completed')
        
        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        if not os.path.exists(file_path):
            return jsonify({"error": "Guide not found"}), 404
            
        with open(file_path, 'r') as f:
            guide_data = json.load(f)
            
        if 'study_schedule' in guide_data and 0 <= session_index < len(guide_data['study_schedule']):
            guide_data['study_schedule'][session_index]['completed'] = completed
            
            with open(file_path, 'w') as f:
                json.dump(guide_data, f)
                
            return jsonify({"success": True, "schedule": guide_data['study_schedule']}), 200
        return jsonify({"error": "Invalid session index"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/guide/<guide_id>/replan', methods=['POST'])
def replan_schedule(guide_id):
    try:
        configure_genai(request)
        data = request.json or {}
        missed_reason = data.get('missed_reason', 'Not specified')

        file_path = os.path.join(DB_PATH, f'{guide_id}.json')
        with open(file_path, 'r') as f:
            guide_data = json.load(f)
            
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        remaining_tasks = [s for s in guide_data.get('study_schedule', []) if not s.get('completed')]
        
        prompt = (
            f"The user has specific study goals: {guide_data.get('goals', 'General mastery')}. "
            "They have NOT completed the following sessions from their previous plan: "
            f"{json.dumps(remaining_tasks)}. "
            f"Reason for missing tasks: '{missed_reason}'. "
            "Please generate a NEW, updated study schedule that helps them catch up. "
            "1. Explain WHY you changed the plan based on their reason (e.g., if busy, make sessions shorter). "
            "2. Adapt the schedule (insert revisions, adjust duration). "
            "Return JSON: {\"study_schedule\": [{\"day_offset\": 1, \"title\": \"...\", \"details\": \"...\", \"duration_minutes\": 30, \"type\": \"learning\"}], \"plan_explanation\": \"...\" }"
        )
        
        response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
        new_plan = json.loads(response.text)
        
        guide_data['study_schedule'] = new_plan['study_schedule']
        guide_data['plan_explanation'] = new_plan.get('plan_explanation', 'Plan updated based on your progress.')
        
        with open(file_path, 'w') as f:
            json.dump(guide_data, f)
            
        return jsonify({"study_schedule": guide_data['study_schedule'], "plan_explanation": guide_data['plan_explanation']}), 200
    except Exception as e:
        logger.error(f"Replan failed: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/motivation', methods=['POST'])
def get_motivation():
    try:
        configure_genai(request)
        data = request.json
        completed_count = data.get('completed_count', 0)
        total_count = data.get('total_count', 0)
        
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        prompt = (
            f"User has completed {completed_count} out of {total_count} study sessions. "
            "Give them a short, punchy, 1-sentence motivational quote or nudge to keep going. "
            "Don't be generic. Be witty if possible."
        )
        response = model.generate_content(prompt)
        return jsonify({"message": response.text.strip()}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=8080)
